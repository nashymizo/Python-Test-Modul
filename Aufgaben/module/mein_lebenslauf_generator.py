"""

    Ein Lebenslaufgenerator

    # Eingabe
    - Daten im JSON Format ablegen
    - (evtl. eine Word Vorlage mit eigenen Stilen)

    # Daten in Word oder PDF aufbereiten
    - Daten (Text)
    - Kopfzeile & Fußzeile (Grafik)
    - Bild/Foto

    # Anforderungen
    - Persönliche Daten
    - Schulischer Werdegang
    - Beruflicher Werdegang (Ausbildung)
    - Kenntnisse und Fähigkeiten

    - Die Daten in eine extra Datei (JSON oder Python)

"""
import json
import time

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls
from docx.shared import Inches
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT

# from pdf.create_simple_pdf import width

meineDaten = {
    "personal": {
        "firstname": "Nico",
        "lastname": "Kehl",
        "birthday": "0*.0*.199*",
        "address": {
            "postal code": "66***",
            "city": "Klei*****",
            "street": "Saargemü*******"
        },
        "Kontakt": {
            "email": "Ni**********",
            "phone": "01***/*70****"
        }
    },

    "education": [
        {
            "school": "Grundschule",
            "degree": "Abschluss",
            "start_date": "1998",
            "end_date": "2002"
        },
        {
            "school": "Weiterführende Schule",
            "degree": "Abschluss",
            "start_date": "2002",
            "end_date": "2006"
        }
    ],

    "experience": [
        {
            "company": "IHK",
            "position": "Prüfer",
            "start_date": "2000",
            "end_date": "2006"
        },
        {
            "company": "Comcave",
            "position": "Umschüler, Fachinformatik AE",
            "start_date": "hier",
            "end_date": "dort"
        }
    ],

    "skills" : [
        "Python",
        "Git",
        "SQL"
    ]
}



def daten_speichern(liste):
    try:
        with open ("Lebenslauf_Daten.json", "w", encoding="utf-8") as file:
            json.dump(liste, file, ensure_ascii=False, indent=4)
            print("Daten wurden erfolgreich gespeichert werden.")

    except Exception as e:
        print(f"Daten konnten nicht gespeichert werden, Fehlermeldung: {e}")


def daten_laden(dateipfad="Lebenslauf_Daten.json"):
    try:
        with open("Lebenslauf_Daten.json", "r", encoding="utf-8") as file:
            daten = json.load(file)
            return daten

    except Exception as e:
        print(f"Datei Upload fehlgeschlagen, Fehlermeldung: {e}")


def kopfzeile_erstellen(doc, personal_daten, bild_pfad):

    tabelle = doc.add_table(rows=1, cols=2)
    tabelle.autofit = False

    tabelle.columns[0].width = Pt(250)  # Linke Spalte schmal
    tabelle.columns[1].width = Pt(330)  # Rechte Spalte breit


    zeile = tabelle.rows[0].cells
    zeile[0].width = Pt(250)
    zeile[1].width = Pt(330)


    p_name = zeile[1].paragraphs[0]
    p_name.text = f"{personal_daten['firstname']} {personal_daten['lastname']}"
    p_name.style = "Heading 1"
    p_name.paragraph_format.space_before = Pt(55)


    if bild_pfad:
        p_bild = zeile[0].paragraphs[0]
        p_bild.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_bild = p_bild.add_run()
        run_bild.add_picture(bild_pfad, width=Inches(1.2), height=Inches(1.7))

    blauen_abschnitt_erstellen(mein_doc, "Persönliche Daten")

    # Tabelle mit 3 Zeilen und 2 Spalten erstellen
    table = doc.add_table(rows=3, cols=2)
    table.autofit = False
    table.columns[0].width = Pt(210)  # Linke Spalte schmal
    table.columns[1].width = Pt(330)  # Rechte Spalte breit

    for row in table.rows:
        row.cells[0].width = Pt(210)
        row.cells[1].width = Pt(330)


    addr = personal_daten['address']
    address_multiline = f"{addr['postal code']} {addr['city']}\n{addr['street']}"

    # Zeile 0: Geburtsdatum
    table.cell(0, 0).paragraphs[0].add_run("Geburtsdatum:").bold = True
    table.cell(0, 1).paragraphs[0].add_run(personal_daten['birthday'])

    # Zeile 1: Adresse
    table.cell(1, 0).paragraphs[0].add_run("Adresse:").bold = True
    table.cell(1, 1).paragraphs[0].add_run(address_multiline)

    # Zeile 2: Kontakt
    table.cell(2, 0).paragraphs[0].add_run("Kontakt:").bold = True
    table.cell(2, 1).paragraphs[0].add_run(f"{personal_daten['Kontakt']['email']} | Tel: {personal_daten['Kontakt']['phone']}")



def bildung_erstellen(doc, bildung_daten):

    blauen_abschnitt_erstellen(doc, "Schulischer Werdegang")

    table = doc.add_table(rows=0, cols=2)
    table.autofit = False


    for eintrag in bildung_daten:

        # Fügt automatisch eine neue Zeile unten an
        row_cells = table.add_row().cells

        # Optional: Spaltenbreiten pro Zeile sichern
        row_cells[0].width = Pt(210)
        row_cells[1].width = Pt(330)

        p_links = row_cells[0].paragraphs[0]
        p_links.paragraph_format.space_after = Pt(20)
        run_zeit = p_links.add_run(f"{eintrag['start_date']} bis {eintrag['end_date']}: ")
        run_zeit.bold = True


        p_rechts = row_cells[1].paragraphs[0]
        p_rechts.paragraph_format.space_after = Pt(20)
        p_rechts.add_run(f"{eintrag['school']} – {eintrag['degree']}")





def beruf_erstellen(doc, beruf_daten):

    blauen_abschnitt_erstellen(doc, "Beruflicher Werdegang")

    table = doc.add_table(rows=0, cols=2)
    table.autofit = False

    for eintrag in beruf_daten:

        row_cells = table.add_row().cells

        row_cells[0].width = Pt(210)
        row_cells[1].width = Pt(330)

        p_links = row_cells[0].paragraphs[0]
        p_links.paragraph_format.space_after = Pt(20)
        run_zeit = p_links.add_run(f"{eintrag['start_date']} bis {eintrag['end_date']}: ")
        run_zeit.bold = True

        p_links = row_cells[1].paragraphs[0]
        p_links.paragraph_format.space_after = Pt(20)
        p_links.add_run(f"{eintrag['company']} – {eintrag['position']}")



    # for eintrag in beruf_daten:
    #     p = doc.add_paragraph()
    #     p.paragraph_format.space_after = Pt(4)
    #
    #     # Run 1: Zeitraum (fett gedruckt)
    #     run_zeit = p.add_run(f"{eintrag['start_date']} bis {eintrag['end_date']}: ")
    #     run_zeit.bold = True
    #
    #     # Run 2: Firma & Position (normaler Text)
    #     p.add_run(f"{eintrag['company']} – {eintrag['position']}")


def skill_erstellen(doc, skill_daten):
    blauen_abschnitt_erstellen(doc, "Skills")


    for eintrag in skill_daten:
        # style="List Bullet" erzeugt die Stichpunkte
        doc.add_paragraph(f"   •  {eintrag}")




def blauen_abschnitt_erstellen(doc, titel):

    p = doc.add_paragraph()

   # Perfekte Abstände definieren:
    p.paragraph_format.space_before = Pt(24)  # Abstand nach oben
    p.paragraph_format.space_after = Pt(17)  # Abstand nach unten zur Linie/Tabelle


    run = p.add_run(titel)
    # zeile formatieren - wie größe, fett, farbe ... in dem beispiel
    run.font.size = Pt(15)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 112, 192)

    # Untere Linie für den Abschnitt (Trennlinie) - copy paste
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(r'<w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                     r'<w:bottom w:val="single" w:sz="6" w:space="4" w:color="CCCCCC"/>'
                     r'</w:pBdr>')
    pPr.append(pBdr)





daten_speichern(meineDaten)
meine_daten = daten_laden()
mein_doc = Document("Vorlage.docx")
kopfzeile_erstellen(mein_doc, meine_daten["personal"], "kakarot1.png")
bildung_erstellen(mein_doc, meine_daten["education"])
beruf_erstellen(mein_doc, meine_daten["experience"])
skill_erstellen(mein_doc, meine_daten["skills"])
mein_doc.save("Lebenslauf.docx")
